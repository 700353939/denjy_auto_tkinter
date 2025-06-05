import os
from tkinter import messagebox, filedialog
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfgen import canvas
from reportlab.pdfbase.ttfonts import TTFont

def export_repair_to_txt(context, repair, car, client):

    file_path = filedialog.asksaveasfilename(
        parent=context.master,
        defaultextension=".txt",
        filetypes=[("txt файлове", "*.txt")],
        title="Запиши ремонта като TXT"
    )

    if not file_path: # If Cancel
        return

    try:
        with open(file_path, "w", encoding="utf-8") as f:
            f.write("\nD E N J Y   A U T O\n")
            f.write(f"\nКлиент: {client.name}, "
                       f"Дата: {repair.repair_date}, "
                       f"Автомобил: {car.registration_number}, "
                       f"Километри: {repair.repair_km}\n")
            f.write(f"Видове ремонти:\n{repair.repairs_type_field}\n")
            f.write(f"Цена: {repair.repair_price:.2f} лв. - платено: {'Да' if repair.is_it_paid else 'Не'}\n")

        messagebox.showinfo("Готово", "Файлът е записан успешно.")

        os.startfile(file_path)

    except Exception as e:
        messagebox.showerror("Грешка", f"Неуспешно записване: {e}")

def export_repair_to_docx(context, repair, car, client):
    docx = Document()
    docx.add_picture("images/denjyauto.gif")
    docx.paragraphs[-1].alignment = 1
    docx.add_heading("Детайли за ремонт", level=1)
    docx.add_paragraph()
    docx.add_paragraph(f"Клиент: {client.name}, "
                       f"Дата: {repair.repair_date}, "
                       f"Автомобил: {car.registration_number}, "
                       f"Километри: {repair.repair_km}")

    docx.add_paragraph(f"Видове ремонти:\n{repair.repairs_type_field}")
    docx.add_paragraph(f"Цена: {repair.repair_price:.2f} лв. - платено: {'Да' if repair.is_it_paid else 'Не'}")

    file_path = filedialog.asksaveasfilename(
        parent=context.master,
        defaultextension=".docx",
        filetypes=[("docx файлове", "*.docx")],
        title="Запиши ремонта като docx"
    )
    docx.save(file_path)

    messagebox.showinfo("Готово", f"Файлът е създаден:\n{file_path}")
    os.startfile(file_path)

def export_repair_to_pdf(context, repair, car, client):

    pdfmetrics.registerFont(TTFont("JetBrainsMono", "fonts/JetBrainsMono-ExtraLight.ttf"))

    file_path = filedialog.asksaveasfilename(
        parent=context.master,
        defaultextension=".pdf",
        filetypes=[("pdf файлове", "*.pdf")],
        title="Запиши ремонта като pdf"
    )

    c = canvas.Canvas(file_path, pagesize=A4)
    width, height = A4
    y = height - 50

    image_path = "images/denjyauto.gif"
    try:
        img_width = 2.0 * inch
        x_image = (width - img_width) / 2
        c.drawImage(image_path, x_image, y - 150, width=img_width, preserveAspectRatio=True, mask='auto')
        y -= 150
    except Exception as e:
        print(f"Не може да се зареди логото: {e}")

    c.setFont("JetBrainsMono", 16)
    c.drawString(50, y, "Детайли за ремонт")
    y -= 40

    c.setFont("JetBrainsMono", 12)
    c.drawString(50, y, f"Клиент: {client.name}, Автомобил: {car.registration_number}")
    y -= 20
    c.drawString(50, y,f"Дата: {repair.repair_date}, Километри: {repair.repair_km}")
    y -= 20
    c.drawString(50, y, f"Видове ремонт:")
    y -= 20
    for line in repair.repairs_type_field.split(", "):
        c.drawString(50, y, line)
        y -= 20

    c.drawString(50, y, f"Цена: {repair.repair_price:.2f} лв. - платено: {'Да' if repair.is_it_paid else 'Не'}")
    y -= 20
    c.showPage()
    c.save()

    messagebox.showinfo("Готово", f"PDF файлът е създаден:\n{file_path}")
    os.startfile(file_path)
